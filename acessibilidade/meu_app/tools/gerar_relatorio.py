from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from io import BytesIO

# Configuração do backend do Matplotlib
import matplotlib
matplotlib.use('Agg')  # Configura o backend para evitar dependências gráficas
import matplotlib.pyplot as plt


def criar_grafico(total, com_alt, sem_alt):
    """
    Cria um gráfico de barras e retorna como um objeto BytesIO.
    """
    plt.figure(figsize=(6, 4))  # Configura o tamanho do gráfico
    plt.bar(['Total', 'Com Alt', 'Sem Alt'], [total, com_alt, sem_alt], color=['blue', 'green', 'red'])
    plt.title('Distribuição das Imagens', fontsize=14)
    plt.xlabel('Categorias', fontsize=12)
    plt.ylabel('Quantidade', fontsize=12)
    plt.tight_layout()

    # Salva o gráfico em memória como um arquivo PNG
    buffer = BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    plt.close()  # Fecha o gráfico para liberar memória
    return buffer


def gerar_relatorio_docx(resultado_analises, nome_arquivo='relatorio_auditoria.docx'):
    """
    Gera e salva o relatório DOCX com os resultados da análise.
    """
    doc = Document()
    doc.add_heading('Relatório de Auditoria de Imagens', level=1)

    for url, detalhes in resultado_analises.items():
        # Adiciona o título da URL analisada
        doc.add_heading(f'URL: {url}', level=2)

        # Adiciona gráfico
        grafico = criar_grafico(
            detalhes['total_imagens'],
            detalhes['imagens_com_alt'],
            detalhes['qtd_imagens_sem_alt']
        )
        doc.add_paragraph("Distribuição das Imagens:").alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adiciona gráfico ao relatório
        grafico_buffer = grafico
        doc.add_picture(grafico_buffer, width=Inches(5))

        # Adiciona tabelas com os detalhes das imagens sem Alt
        doc.add_heading('Detalhes das Imagens Sem Texto Alternativo', level=3)
        for idx, imagem in enumerate(detalhes['detalhes_imagens_sem_alt'], start=1):
            # Adicionar uma tabela para cada imagem
            tabela = doc.add_table(rows=3, cols=2)
            tabela.style = 'Table Grid'
            tabela.cell(0, 0).text = 'Imagem'
            tabela.cell(0, 1).text = 'Texto Alternativo'

            # Formatar cabeçalhos
            for cell in tabela.rows[0].cells:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Adicionar URL e imagem
            tabela.cell(1, 0).text = imagem['img_url']
            img_path = os.path.join('img', imagem['img_url'].split('/')[-1])
            if os.path.exists(img_path):
                tabela.cell(2, 0).paragraphs[0].add_run().add_picture(img_path, width=Inches(2))
            else:
                tabela.cell(2, 0).text = "Imagem não encontrada."

            # Adicionar texto alternativo
            tabela.cell(1, 1).merge(tabela.cell(2, 1))
            tabela.cell(1, 1).text = imagem.get('texto_alternativo_sugerido', 'Texto alternativo não disponível')
            tabela.cell(1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            doc.add_paragraph()  # Espaçamento entre tabelas

    # Salvar o documento
    doc.save(nome_arquivo)
    print(f"Relatório salvo como {nome_arquivo}")


def salvar_relatorio(resultado_analises):
    """
    Envolve a função gerar_relatorio_docx para salvar o relatório e retornar o nome do arquivo.
    """
    nome_arquivo_docx = "relatorio_auditoria.docx"
    gerar_relatorio_docx(resultado_analises, nome_arquivo_docx)
    return nome_arquivo_docx
